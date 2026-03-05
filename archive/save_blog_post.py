"""Save the Medium blog post as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "output",
    "Funds_Flow_Indexer_Medium_Post.docx"
)

TITLE = "I Built a Deal Closing Auditor That Never Sleeps — and It's a Dream for Both Auditors and Preparers"
SUBTITLE = "How an agentic workflow turns a chaotic stack of PDFs and a funds flow spreadsheet into a clean, documented audit trail in minutes."

SECTIONS = [
    {
        "heading": "The Problem Nobody Talks About",
        "body": (
            "Every private equity deal closes with the same messy ritual.\n\n"
            "The finance team has a funds flow statement — a spreadsheet listing every dollar leaving (or entering) "
            "the firm at closing: legal fees, due diligence invoices, banking advisory, filing fees, the actual purchase price. "
            "Dozens of line items. Millions of dollars.\n\n"
            "Somewhere in a shared drive there's a folder of support documents: invoices, wire confirmations, engagement letters, "
            "email threads with PDF attachments. These are supposed to prove every number on the funds flow.\n\n"
            "The problem? Nobody mapped them together.\n\n"
            "The auditor arrives and asks: \"Can you show me support for each line item?\" The preparer spends the next two days "
            "manually hunting through folders, renaming files, cross-referencing amounts, building a workpaper from scratch — "
            "work that adds zero analytical value but consumes enormous time.\n\n"
            "I built a tool that does all of it automatically. And it turns out, getting this right required something more than "
            "deterministic rules."
        ),
    },
    {
        "heading": "What the Tool Does",
        "body": (
            "At its core, the Funds Flow Audit Tool takes two inputs:\n\n"
            "  1. A client funds flow Excel workbook\n"
            "  2. A folder of support documents (PDFs, email invoices)\n\n"
            "And produces three outputs:\n\n"
            "  1. An annotated workpaper — the client's Excel file with GL account codes, fund allocation formulas, "
            "and cross-reference numbers added in-place, auditor-style\n"
            "  2. FF-numbered documents — every matched invoice renamed and filed as FF01_HarringtonCross_Invoice.pdf, "
            "FF02_..., etc., exactly how a workpaper binder should look\n"
            "  3. A journal entry tab embedded in the same workbook — debits and credits formula-linked back to the "
            "annotated columns, split by fund, with balance checks\n\n"
            "The entire audit trail is self-contained in a single Excel file with formula transparency. An auditor can trace "
            "every number from the journal entry back to the funds flow amount back to the source document — without leaving "
            "the workbook."
        ),
    },
    {
        "heading": "The Architecture: Two Layers That Shouldn't Be Confused",
        "body": (
            "The tool is built as an agentic workflow — a term that gets thrown around loosely, but here it has a precise meaning.\n\n"
            "Layer 1 — Deterministic Python scripts. These handle the I/O work that has no ambiguity: extract rows from Excel tabs, "
            "extract text from PDFs, write the annotated workbook, rename and file documents, compute formula cells. These scripts "
            "are fast, repeatable, and require no intelligence. They do exactly what they're told, every time.\n\n"
            "Layer 2 — Claude Code as the judgment layer. This is where the interesting work happens. A skill (a structured prompt, "
            "essentially a set of instructions) tells Claude what to do and in what order. But Claude's own reasoning is what "
            "actually executes the judgment: matching line items to documents, classifying GL accounts, resolving ambiguity, "
            "handling edge cases. The skill is the workflow. The intelligence is Claude's.\n\n"
            "This is an important distinction. A lot of \"AI-powered\" tools hardcode the intelligence into Python rules — keyword "
            "lists, regex patterns, lookup tables. I deliberately removed all of that. The chart of accounts is a pure reference "
            "list (codes and names, nothing more). The GL classification happens entirely in the LLM's reasoning, not in code."
        ),
    },
    {
        "heading": "Where This Gets Interesting: Edge Cases",
        "body": (
            "A real funds flow isn't clean. The demo deal I built to test this has exactly the kinds of messiness you'd find in a "
            "real closing:\n\n"
            "Cumulative billing. A law firm sends Invoice 001 for $350,000 in January, then Invoice 002 for $750,000 in July — "
            "but Invoice 002 is cumulative, meaning it represents the total engagement to date, not just the July work. The funds "
            "flow has two line items ($350K and $400K incremental). There's only one final document that covers the combined total. "
            "The LLM has to recognize that these two line items share a vendor, that the document amount is larger than either "
            "individual line, and that the two together reconcile to the document — and assign a CUMULATIVE status accordingly.\n\n"
            "Partial invoices. The HSR filing fee line item is $30,000. The government filing fee receipt is only $25,000. The "
            "remaining $5,000 was billed separately for an SEC pre-notification. The LLM assigns PARTIAL, links the document it "
            "found, and flags the unreconciled $5,000 — rather than silently calling it matched.\n\n"
            "Missing support. Two line items have no documents at all: a strategy consulting firm and a tax advisory engagement. "
            "Rather than failing, the LLM assigns MISSING, still classifies the GL account from context, and includes those items "
            "in the summary with clear flags.\n\n"
            "Email invoices. One vendor sends invoices as email PDFs. The LLM recognizes the email format from context and sets "
            "document_type: \"email\" rather than \"invoice\", which matters for how the workpaper documents the source.\n\n"
            "Orphan documents. Some PDFs in the folder don't correspond to any line item — a duplicate engagement letter, a general "
            "research report. These get routed to an UNMATCHED/ folder automatically, so the auditor knows they exist without "
            "cluttering the main index.\n\n"
            "None of these edge cases require code changes. They're handled by reasoning over the actual content of the documents "
            "and the funds flow, not by rules written in advance."
        ),
    },
    {
        "heading": "GL Classification Without Rules",
        "body": (
            "One design decision I'm particularly happy with: there are no keyword-matching rules anywhere in the codebase.\n\n"
            "Early versions had an ACCOUNT_RULES dictionary in Python — \"if description contains 'legal', assign 7010; if it "
            "contains 'advisory', assign 7020.\" This broke almost immediately. A \"Tax Structuring & Transaction Advisory\" "
            "engagement hit the \"advisory\" rule before it reached the \"tax\" rule and got classified wrong. The fix kept being "
            "more rules, more edge cases, more maintenance.\n\n"
            "The solution was to remove the rules entirely. The chart of accounts is now a reference list:\n\n"
            "  7010 — Transaction Costs — Legal Fees\n"
            "  7020 — Transaction Costs — Banking & Advisory Fees\n"
            "  7030 — Transaction Costs — Financial Due Diligence\n"
            "  ...\n"
            "  1500 — Investment in Portfolio Companies\n\n"
            "And the skill tells Claude: read the description, the vendor name, the document type, and the context. Reason about "
            "what the expense actually is. Don't match keywords mechanically.\n\n"
            "\"SPA execution\" -> Legal Fees. \"Fairness opinion\" -> Banking & Advisory. \"QofE analysis\" -> Financial Due "
            "Diligence. \"HSR filing fee\" -> Regulatory & Filing Fees. The LLM gets these right because it understands what the "
            "words mean in a deal context — not because a rule told it to."
        ),
    },
    {
        "heading": "The Output an Auditor Actually Wants",
        "body": (
            "The annotated workbook opens to the client's original funds flow, completely intact, with new columns added to the "
            "right: GL code, GL account name, Fund I amount (formula), Fund II amount (formula), and a check column confirming "
            "the split adds back to the total.\n\n"
            "A separate Journal Entry tab has two complete journal entry sections — one for Fund I, one for Fund II — each with "
            "debit lines referencing the annotated cells via Excel formulas, credit lines to Wire Payable and Accrued Liabilities, "
            "and a balance check row. Change one number in the funds flow and the journal entry updates automatically.\n\n"
            "The documents folder becomes documents_indexed/ — every matched PDF renamed with its FF reference number, every "
            "unmatched document quarantined in an UNMATCHED/ subfolder.\n\n"
            "A summary report prints to the terminal:\n\n"
            "  =====================================================================\n"
            "  FUNDS FLOW INDEX — Project Meridian\n"
            "  Closing: 2026-07-15   Role: BUYER\n"
            "  =====================================================================\n"
            "  FF#   Description                          Amount         Status\n"
            "  FF01  Harrington & Cross — Retainer        $350,000       CUMULATIVE\n"
            "  FF02  Harrington & Cross — Closing         $400,000       CUMULATIVE\n"
            "  FF03  Meridian Advisory Group — QofE       $280,000       MATCHED\n"
            "  FF04  Pinnacle Capital — Advisory          $2,750,000     MATCHED\n"
            "  —     Vertex Strategy Group                $300,000       MISSING\n"
            "  =====================================================================\n"
            "  Matched: 6  |  Partial: 1  |  Cumulative: 2  |  Missing: 2\n"
            "  Supported:  $4,245,000   |   Unsupported: $450,000\n"
            "  Unmatched documents: 2 (see UNMATCHED/)\n"
            "  ====================================================================="
        ),
    },
    {
        "heading": "Why This Matters",
        "body": (
            "The work of matching invoices to a funds flow is not intellectually difficult — but it is painstaking, error-prone, "
            "and time-consuming when done manually. It's exactly the kind of work where humans make mistakes not from lack of "
            "skill but from fatigue and volume.\n\n"
            "What makes an LLM the right tool here isn't that it's smarter than a senior associate. It's that it can hold the "
            "full context of every line item and every document simultaneously, reason about relationships across all of them, "
            "apply consistent judgment without getting tired, and document its reasoning in a structured output that feeds "
            "directly into a clean audit trail.\n\n"
            "The preparer doesn't spend two days building a binder. The auditor gets a self-documented workpaper with formula "
            "traceability on day one. Everyone gets to spend their time on the work that actually requires judgment — not the indexing."
        ),
    },
]

FOOTER = (
    "The Funds Flow Audit Tool is a local tool built with Python, openpyxl, pdfplumber, and Claude. "
    "It runs entirely on your machine and calls no external APIs during document processing. "
    "The LLM judgment layer runs through Claude Code."
)


def build_doc():
    doc = Document()

    # Title
    t = doc.add_heading(TITLE, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in t.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Subtitle / deck
    sub = doc.add_paragraph(SUBTITLE)
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in sub.runs:
        run.font.size = Pt(13)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    for section in SECTIONS:
        h = doc.add_heading(section["heading"], level=2)
        for run in h.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        body = doc.add_paragraph(section["body"])
        body.paragraph_format.space_after = Pt(10)
        for run in body.runs:
            run.font.size = Pt(11)

        doc.add_paragraph()  # spacer between sections

    # Footer / italics note
    doc.add_paragraph("---")
    footer = doc.add_paragraph(FOOTER)
    for run in footer.runs:
        run.font.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_doc()
